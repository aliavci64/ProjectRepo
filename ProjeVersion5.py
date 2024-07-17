import os
import subprocess
from docx import Document
from docx.shared import RGBColor
import re
import sys

def get_git_diff(commit1, commit2, repo_path, file_name):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', '-U0', commit1, commit2, '--', file_name],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8'
        )
        return result.stdout
    except Exception as e:
        return str(e)

def extract_function_names(diff):
    function_names = set()
    pattern = re.compile(r'^@@.*@@\s*(def|function)\s*([a-zA-Z_][a-zA-Z0-9_]*)')
    for line in diff.splitlines():
        match = pattern.match(line)
        if match:
            function_names.add(match.group(2))
    return function_names

def write_diff_to_word(diff, output_file, file_status=None, file_name=None):
    try:
        doc = Document()
        doc.add_heading(f'{file_name} Dosyasındaki Farklılıklar:', level=1)
        
        if file_status:
            if file_status == 'A':
                doc.add_paragraph(f"Dosya Eklendi: {file_name}").runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif file_status == 'D':
                doc.add_paragraph(f"Dosya Silindi: {file_name}").runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            lines = diff.splitlines()
            difference_count = sum(1 for line in lines if line.startswith('-') or line.startswith('+'))
            function_names = extract_function_names(diff)

            doc.add_paragraph(f"Farklılık Sayısı: {difference_count}")
            doc.add_paragraph(f"Farklılık Olan Fonksiyonlar: {', '.join(function_names)}")

            old_line_num = 0
            new_line_num = 0

            for line in lines:
                if line.startswith('@@'):
                    parts = line.split(' ')
                    old_line_num = int(parts[1].split(',')[0][1:])
                    new_line_num = int(parts[2].split(',')[0][1:])
                elif line.startswith('-'):
                    run = doc.add_paragraph().add_run(f"{old_line_num}: {line[1:]}")
                    font = run.font
                    font.color.rgb = RGBColor(0, 128, 0)  # Yeşil renk (eklenen satırlar)
                    old_line_num += 1
                elif line.startswith('+'):
                    run = doc.add_paragraph().add_run(f"{new_line_num}: {line[1:]}")
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk (silinen satırlar)
                    new_line_num += 1

        doc.save(output_file)
    except Exception as e:
        print(f"Word belgesine yazarken hata oluştu: {e}")

def get_changed_files(commit1, commit2, repo_path):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', '--name-status', commit1, commit2],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8'
        )
        return result.stdout.splitlines()
    except Exception as e:
        return str(e)

def get_commit_diff(commit, repo_path):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff-tree', '--no-commit-id', '--name-status', '-r', commit],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8'
        )
        return result.stdout.splitlines()
    except Exception as e:
        return str(e)

def convert_path(path):
    return path.replace('\\', '/')

def create_unique_output_dir(base_dir):
    counter = 1
    new_dir = base_dir
    while os.path.exists(new_dir):
        new_dir = f"{base_dir} ({counter})"
        counter += 1
    os.makedirs(new_dir)
    return new_dir

if __name__ == "__main__":
    if len(sys.argv) < 4 or len(sys.argv) > 5:
        print("Usage: python diff_to_word.py <commit1> [<commit2>] <repo_path> <output_dir>")
        sys.exit(1)
    
    commit1 = sys.argv[1]
    if len(sys.argv) == 5:
        commit2 = sys.argv[2]
        repo_path = convert_path(sys.argv[3])
        output_dir = convert_path(sys.argv[4])
    else:
        commit2 = None
        repo_path = convert_path(sys.argv[2])
        output_dir = convert_path(sys.argv[3])

    print(f"Converted repo_path: {repo_path}")
    print(f"Converted output_dir: {output_dir}")

    differences_dir = create_unique_output_dir(os.path.join(repo_path, "Differences"))

    if commit2:
        changed_files = get_changed_files(commit1, commit2, repo_path)
    else:
        changed_files = get_commit_diff(commit1, repo_path)

    for line in changed_files:
        status, file_name = line.split(maxsplit=1)
        output_file = os.path.join(differences_dir, f"{os.path.basename(file_name)}_diff.docx")
        
        if status == 'A':
            write_diff_to_word('', output_file, 'A', file_name)
        elif status == 'D':
            write_diff_to_word('', output_file, 'D', file_name)
        else:
            diff = get_git_diff(commit1, commit2 if commit2 else commit1 + '^', repo_path, file_name)
            write_diff_to_word(diff, output_file, file_name=file_name)

    print(f"Farklılıklar {differences_dir} klasörüne ayrı Word dosyaları olarak yazıldı.")
