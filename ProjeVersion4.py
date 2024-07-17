import os
import subprocess
from docx import Document
from docx.shared import RGBColor
import re
import sys

def get_git_diff(commit1, commit2, repo_path, file_name):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', '-U0', commit1, commit2, '--', file_name],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        return result.stdout
    except Exception as e:
        return str(e)

def extract_function_names(diff):
    function_names = set()
    pattern = re.compile(r'^@@.*@@\s*(def|function)\s*([a-zA-Z_][a-zA-Z0-9_]*)')
    for line in diff.splitlines():
        match = pattern.match(line)
        if match:
            function_names.add(match.group(2))
    return function_names

def write_diff_to_word(diff, output_file, file_status=None, file_name=None):
    try:
        doc = Document()
        doc.add_heading(f'File Differences for {file_name}:', level=1)
        
        if file_status:
            if file_status == 'A':
                doc.add_paragraph(f"File Added: {file_name}").runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif file_status == 'D':
                doc.add_paragraph(f"File Deleted: {file_name}").runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            lines = diff.splitlines()
            difference_count = sum(1 for line in lines if line.startswith('-') or line.startswith('+'))
            function_names = extract_function_names(diff)

            doc.add_paragraph(f"Number of differences: {difference_count}")
            doc.add_paragraph(f"Functions with differences: {', '.join(function_names)}")

            old_line_num = 0
            new_line_num = 0

            for line in lines:
                if line.startswith('@@'):
                    parts = line.split(' ')
                    old_line_num = int(parts[1].split(',')[0][1:])
                    new_line_num = int(parts[2].split(',')[0][1:])
                elif line.startswith('-'):
                    run = doc.add_paragraph().add_run(f"{old_line_num}: {line[1:]}")
                    font = run.font
                    font.color.rgb = RGBColor(0, 128, 0)  # Yeşil renk (eklenen satırlar)
                    old_line_num += 1
                elif line.startswith('+'):
                    run = doc.add_paragraph().add_run(f"{new_line_num}: {line[1:]}")
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk (silinen satırlar)
                    new_line_num += 1

        doc.save(output_file)
    except Exception as e:
        print(f"Error writing to Word document: {e}")

def get_changed_files(commit1, commit2, repo_path):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', '--name-status', commit1, commit2],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        return result.stdout.splitlines()
    except Exception as e:
        return str(e)

def convert_path(path):
    return path.replace('\\', '/')

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python diff_to_word.py <commit1> <commit2> <repo_path> <output_dir>")
        sys.exit(1)
    
    commit1 = sys.argv[1]
    commit2 = sys.argv[2]
    repo_path = convert_path(sys.argv[3])
    output_dir = convert_path(sys.argv[4])

    print(f"Converted repo_path: {repo_path}")
    print(f"Converted output_dir: {output_dir}")

    differences_dir = os.path.join(repo_path, "Differences")
    os.makedirs(differences_dir, exist_ok=True)

    changed_files = get_changed_files(commit1, commit2, repo_path)

    for line in changed_files:
        status, file_name = line.split(maxsplit=1)
        output_file = os.path.join(differences_dir, f"{os.path.basename(file_name)}_diff.docx")
        
        if status == 'A':
            write_diff_to_word('', output_file, 'A', file_name)
        elif status == 'D':
            write_diff_to_word('', output_file, 'D', file_name)
        else:
            diff = get_git_diff(commit1, commit2, repo_path, file_name)
            write_diff_to_word(diff, output_file, file_name=file_name)

    print(f"Differences written to separate Word files in {differences_dir}")
