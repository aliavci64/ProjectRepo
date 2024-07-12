import os
import subprocess
from docx import Document
from docx.shared import RGBColor
import sys

def get_git_diff(commit1, commit2, repo_path, file_name):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', commit1, commit2, '--', file_name],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        return result.stdout
    
    except Exception as e:
        return str(e)

def write_diff_to_word(diff, output_file, file_status=None, file_name=None):
    try:
        doc = Document()
        doc.add_heading('File Differences:', level=1)
        
        if file_status:
            if file_status == 'A':
                doc.add_paragraph(f"File Added: {file_name}").runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif file_status == 'D':
                doc.add_paragraph(f"File Deleted: {file_name}").runs[0].font.color.rgb = RGBColor(255, 0, 0)
        else:
            for line in diff.splitlines():
                if line.startswith('-'):
                    run = doc.add_paragraph().add_run(line[1:])
                    font = run.font
                    font.color.rgb = RGBColor(0, 128, 0)  # Yeşil renk (eklenen satırlar)
                    
                elif line.startswith('+'):
                    run = doc.add_paragraph().add_run(line[1:])
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk (silinen satırlar)
                else:
                    doc.add_paragraph(line)

        doc.save(output_file)
    except Exception as e:
        print(f"Error writing to Word document: {e}")

def get_changed_files(commit1, commit2, repo_path):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', '--name-status', commit1, commit2],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        return result.stdout.splitlines()
    
    except Exception as e:
        return str(e)

def convert_path(path):
    return path.replace('\\', '/')

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python diff_to_word.py <commit1> <commit2> <repo_path> <output_dir>")
        sys.exit(1)
    
    commit1 = sys.argv[1]
    commit2 = sys.argv[2]
    repo_path = convert_path(sys.argv[3])
    output_dir = convert_path(sys.argv[4])

    print(f"Converted repo_path: {repo_path}")
    print(f"Converted output_dir: {output_dir}")

    changed_files = get_changed_files(commit1, commit2, repo_path)

    for line in changed_files:
        status, file_name = line.split(maxsplit=1)
        output_file = os.path.join(output_dir, f"{os.path.basename(file_name)}_diff.docx")
        
        if status == 'A':
            write_diff_to_word('', output_file, 'A', file_name)
        elif status == 'D':
            write_diff_to_word('', output_file, 'D', file_name)
        else:
            diff = get_git_diff(commit1, commit2, repo_path, file_name)
            write_diff_to_word(diff, output_file)

    print(f"Differences written to separate Word files in {output_dir}")
