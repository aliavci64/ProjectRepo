import os
import subprocess
from docx import Document
from docx.shared import RGBColor
import sys 

def get_git_diff(commit1, commit2, repo_path):
    try:
        result = subprocess.run(
            ['git', '-C', repo_path, 'diff', commit1, commit2],
            stdout = subprocess.PIPE, stderr = subprocess.PIPE, text = True
        )
        return result.stdout
    
    except Exception as e:
        return str(e)
    

def write_diff_to_word(diff, output_file):
    try:
        doc = Document()
        doc.add_heading('File Differences:', level = 1)
        for line in diff.splitlines():
            if line.startswith('-'):
                run = doc.add_paragraph().add_run(line[1:])
                font = run.font
                font.color.rgb = RGBColor(0, 128, 0)  # Yeşil renk (eklenen satırlar)
            elif line.startswith('+'):
                run = doc.add_paragraph().add_run(line[1:])
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk (silinen satırlar)
            else:
                doc.add_paragraph(line)

        # Dosya değişikliklerini ekleyecek fonksiyon
        def add_file_change(status, file_name):
            if status == 'A':
                color = RGBColor(0, 128, 0)  # Yeşil renk (eklenen dosya)
                doc.add_paragraph(f"Added: {file_name}", style='List Bullet').runs[0].font.color.rgb = color
            elif status == 'D':
                color = RGBColor(255, 0, 0)  # Kırmızı renk (silinen dosya)
                doc.add_paragraph(f"Deleted: {file_name}", style='List Bullet').runs[0].font.color.rgb = color
            else:
                doc.add_paragraph(f"Modified: {file_name}", style='List Bullet')

        for line in diff.splitlines():
            if line.strip():  # Satır boş değilse işlem yap
                status, file_name = line.split(maxsplit=1)
                add_file_change(status, file_name.strip())

        doc.save(output_file)
    except Exception as e:
        print(f"Error writing to Word document: {e}")

def convert_path(path):
    return path.replace('\\', '/')

if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python diff_to_word.py <commit1> <commit2> <repo_path> <output_file>")
        sys.exit(1)
    
    commit1 = sys.argv[1]
    commit2 = sys.argv[2]
    repo_path = convert_path(sys.argv[3])
    output_file = convert_path(sys.argv[4])

    print(f"Converted repo_path: {repo_path}")
    print(f"Converted output_file: {output_file}")

    diff = get_git_diff(commit1, commit2, repo_path)
    write_diff_to_word(diff, output_file)

    print(f"Diff writen to {output_file}")






